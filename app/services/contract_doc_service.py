# -*- coding: utf-8 -*-
"""
合同文档生成服务
负责合同 Word 文档的模板填充、生成和下载
"""

import os
import re
from datetime import datetime
from pathlib import Path
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches, Cm
from flask import current_app
from utils.database import DBConnection





def amount_to_chinese(amount):
    """
    将金额转换为人民币大写格式
    
    Args:
        amount: 数字金额
        
    Returns:
        str: 人民币大写金额
    """
    if amount is None or amount == 0:
        return '零元整'
    
    # 处理负数
    is_negative = False
    if amount < 0:
        is_negative = True
        amount = abs(amount)
    
    # 人民币大写字符
    chinese_num = ['零', '壹', '贰', '叁', '肆', '伍', '陆', '柒', '捌', '玖']
    chinese_unit = ['', '拾', '佰', '仟']
    chinese_section = ['', '万', '亿', '兆']
    chinese_yuan = '元'
    chinese_jiao = ['角', '分']
    chinese_integer = '整'
    
    # 分离整数和小数部分
    integer_part = int(amount)
    decimal_part = round((amount - integer_part) * 100)
    
    # 转换整数部分
    def convert_integer(n):
        if n == 0:
            return '零'
        
        result = ''
        section_idx = 0
        need_zero = False
        
        while n > 0:
            section = n % 10000
            section_str = ''
            
            # 处理当前 section 内的 4 位
            for i in range(4):
                digit = section % 10
                if digit == 0:
                    # 当前位为 0，标记可能需要添加零
                    if section_str and not section_str.startswith('零'):
                        section_str = '零' + section_str
                else:
                    section_str = chinese_num[digit] + chinese_unit[i] + section_str
                section = section // 10
            
            # 去除开头的零
            section_str = section_str.lstrip('零')
            # 去除末尾的零
            section_str = section_str.rstrip('零')
            
            if section_str:
                section_str += chinese_section[section_idx]
                # 如果前面已有结果，且需要零
                if need_zero and result:
                    section_str = '零' + section_str
                result = section_str + result
                need_zero = False
            elif section_idx > 0 and result:
                # 当前 section 为 0，但前面有结果，标记需要零
                need_zero = True
            
            n = n // 10000
            section_idx += 1
        
        # 去除连续的零
        result = result.replace('零零', '零')
        
        # 特殊处理：万位和个位之间需要零（如 10001 -> 一万零一）
        # 只有当"万"后面直接跟"壹元整"时才添加零
        if section_idx > 1 and len(result) > 1:
            # 只处理"万"后面直接跟个位 + 元整的情况
            result = result.replace('万壹元整', '万零壹元整')
            result = result.replace('万贰元整', '万零贰元整')
            result = result.replace('万叁元整', '万零叁元整')
            result = result.replace('万肆元整', '万零肆元整')
            result = result.replace('万伍元整', '万零伍元整')
            result = result.replace('万陆元整', '万零陆元整')
            result = result.replace('万柒元整', '万零柒元整')
            result = result.replace('万捌元整', '万零捌元整')
            result = result.replace('万玖元整', '万零玖元整')
        
        return result
    
    # 转换整数部分
    integer_str = convert_integer(integer_part)
    if integer_str:
        result = integer_str + chinese_yuan
    else:
        result = '零' + chinese_yuan
    
    # 转换小数部分
    jiao = decimal_part // 10
    fen = decimal_part % 10
    
    if jiao > 0 or fen > 0:
        if integer_part == 0 and jiao > 0:
            # 整数部分为 0 且有角，可以省略"零元"
            result = chinese_num[jiao] + chinese_jiao[0]
            if fen > 0:
                result += chinese_num[fen] + chinese_jiao[1]
            else:
                result += chinese_integer
        elif jiao > 0:
            result += chinese_num[jiao] + chinese_jiao[0]
            if fen > 0:
                result += chinese_num[fen] + chinese_jiao[1]
            else:
                result += chinese_integer
        else:
            result += '零' + chinese_jiao[0] + chinese_num[fen] + chinese_jiao[1]
    else:
        result += chinese_integer
    
    # 处理负数
    if is_negative:
        result = '负' + result
    
    return result


class ContractDocService:
    """合同文档服务类"""
    
    # 允许的文件扩展名
    ALLOWED_EXTENSIONS = {'docx'}
    
    def __init__(self, app=None):
        """初始化服务"""
        self.app = app
        if app is not None:
            self.init_app(app)
    
    def init_app(self, app):
        """初始化 Flask 应用"""
        # 确保生成目录存在
        generated_docs_dir = Path(app.root_path) / 'generated_docs'
        generated_docs_dir.mkdir(parents=True, exist_ok=True)
        
        # 确保模板目录存在
        template_dir = Path(app.root_path) / 'templates'
        template_dir.mkdir(parents=True, exist_ok=True)
    
    @property
    def generated_docs_dir(self):
        """获取生成文档目录"""
        return Path(current_app.root_path) / 'generated_docs'
    
    @property
    def template_dir(self):
        """获取模板目录"""
        return Path(current_app.root_path) / 'templates'
    
    def allowed_file(self, filename):
        """检查文件扩展名是否允许"""
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in self.ALLOWED_EXTENSIONS
    
    def sanitize_filename(self, filename):
        """清理文件名，防止路径遍历攻击"""
        # 移除路径分隔符
        filename = filename.replace('/', '').replace('\\', '')
        # 移除 ..
        filename = filename.replace('..', '')
        # 只保留字母、数字、中文、下划线、中划线和点
        filename = re.sub(r'[^\w\u4e00-\u9fff\-_.]', '', filename)
        return filename
    
    def _insert_plot_images(self, file_path, plots_data):
        """
        在文档末尾插入地块图片
        
        Args:
            file_path: Word 文档路径
            plots_data: 地块数据列表（包含 image_path）
        """
        try:
            from docx import Document
            from docx.shared import Cm
            import time
            
            # 等待文件写入完成（关键）
            time.sleep(0.2)
            
            # 加载已生成的文档
            doc = Document(file_path)
            
            # 强制分页
            doc.add_page_break()
            doc.add_heading('租赁地块图示', level=2)
            
            # 记录已插入的图片数量
            inserted_count = 0
            
            # 遍历每个地块，插入对应的图片
            for plot in plots_data:
                image_path = plot.get('image_path', '')
                
                if not image_path:
                    continue
                
                # 👉 正确路径（关键修复点）
                # image_path 格式：/uploads/plot/xxx.png
                # 文件实际位置：项目根目录/uploads/plot/xxx.png
                # current_app.root_path 是 app 目录，需要向上一级
                project_root = Path(current_app.root_path).parent
                abs_path = os.path.join(project_root, image_path.lstrip('/'))
                
                current_app.logger.info(f'图片路径：{abs_path}')
                
                # 检查文件是否存在
                if os.path.exists(abs_path):
                    try:
                        # 添加地块编号标题
                        doc.add_heading(
                            f'地块编号：{plot["plot_number"]}',
                            level=3
                        )
                        
                        # 插入图片，设置宽度为 14cm，高度自动
                        doc.add_picture(abs_path, width=Cm(14))
                        
                        # 添加段落间距
                        doc.add_paragraph()
                        
                        inserted_count += 1
                        
                    except Exception as e:
                        current_app.logger.error(
                            f'插入图片失败：{plot["plot_number"]}, 错误：{str(e)}'
                        )
                else:
                    current_app.logger.warning(f'图片不存在：{abs_path}')
            
            # 如果没有图片，添加提示
            if inserted_count == 0:
                doc.add_paragraph('暂无地块图片')
            
            # 保存修改后的文档
            doc.save(file_path)
            
        except Exception as e:
            current_app.logger.error(f'插入地块图片失败：{str(e)}')
    
    def get_contract_data(self, contract_id):
        """
        获取合同完整数据
        
        Args:
            contract_id: 合同 ID
            
        Returns:
            dict: 合同数据字典
        """
        # 使用原生 SQL 查询，避免模型依赖
        with DBConnection() as conn:
            cursor = conn.cursor()
            
            # 获取合同信息
            cursor.execute("""
                SELECT c.ContractID, c.ContractNumber, c.ContractName, c.MerchantID,
                       c.ContractPeriod, c.StartDate, c.EndDate, c.ContractAmount,
                       c.AmountReduction, c.ActualAmount, c.PaymentMethod, c.Status,
                       c.Description
                FROM Contract c
                WHERE c.ContractID = ?
            """, (contract_id,))
            
            contract_row = cursor.fetchone()
            if not contract_row:
                return None
            
            # 获取商户信息
            cursor.execute("""
                SELECT MerchantName, LegalPerson, ContactPerson, Phone, Address
                FROM Merchant
                WHERE MerchantID = ?
            """, (contract_row.MerchantID,))
            
            merchant_row = cursor.fetchone()
            
            # 获取地块信息（包含图片路径）
            cursor.execute("""
                SELECT p.PlotID, p.PlotNumber, p.PlotName, p.Area, p.UnitPrice,
                       cp.MonthlyPrice, p.YearlyRent, p.ImagePath
                FROM Plot p
                INNER JOIN ContractPlot cp ON p.PlotID = cp.PlotID
                WHERE cp.ContractID = ?
                ORDER BY p.PlotNumber
            """, (contract_id,))
            
            plot_rows = cursor.fetchall()
        
        # 构建地块列表（包含图片路径）
        plots_data = []
        for idx, plot in enumerate(plot_rows, 1):
            plots_data.append({
                'index': idx,
                'plot_number': plot.PlotNumber,
                'plot_name': plot.PlotName or '',
                'area': float(plot.Area) if plot.Area else 0,
                'unit_price': float(plot.UnitPrice) if plot.UnitPrice else 0,
                'monthly_rent': float(plot.MonthlyPrice) if plot.MonthlyPrice else 0,
                'yearly_rent': float(plot.YearlyRent) if plot.YearlyRent else 0,
                'image_path': plot.ImagePath or '',
            })
        
        # 构建合同数据
        contract_data = {
            'contract_id': contract_row.ContractID,
            'contract_number': contract_row.ContractNumber,
            'contract_name': contract_row.ContractName,
            'merchant_name': merchant_row.MerchantName if merchant_row else '',
            'legal_person': merchant_row.LegalPerson if merchant_row else '',
            'contact_person': merchant_row.ContactPerson if merchant_row else '',
            'phone': merchant_row.Phone if merchant_row else '',
            'address': merchant_row.Address if merchant_row else '',
            'contract_period': contract_row.ContractPeriod or '',
            'start_date': contract_row.StartDate.strftime('%Y年%m月%d日') if contract_row.StartDate else '',
            'end_date': contract_row.EndDate.strftime('%Y年%m月%d日') if contract_row.EndDate else '',
            'contract_amount': f'¥{float(contract_row.ContractAmount):,.2f}' if contract_row.ContractAmount else '¥0.00',
            'amount_reduction': f'¥{float(contract_row.AmountReduction):,.2f}' if contract_row.AmountReduction else '¥0.00',
            'actual_amount': f'¥{float(contract_row.ActualAmount):,.2f}' if contract_row.ActualAmount else '¥0.00',
            'rmb_amount': amount_to_chinese(float(contract_row.ActualAmount)) if contract_row.ActualAmount else '零元整',
            'payment_method': contract_row.PaymentMethod or '银行转账',
            'status': contract_row.Status or '有效',
            'description': contract_row.Description or '',
            'plots': plots_data,
            'total_plots': len(plots_data),
            'generate_time': datetime.now().strftime('%Y年%m月%d日 %H:%M:%S'),
        }
        
        return contract_data
    
    def generate_contract_doc(self, contract_id, template_name='contract_template.docx'):
        """
        生成合同文档（包含地块图片）
        
        Args:
            contract_id: 合同 ID
            template_name: 模板文件名
            
        Returns:
            dict: {
                'success': bool,
                'file_path': str,
                'file_name': str,
                'message': str
            }
        """
        try:
            # 获取合同数据
            contract_data = self.get_contract_data(contract_id)
            if not contract_data:
                return {
                    'success': False,
                    'message': '合同不存在'
                }
            
            # 验证模板文件
            template_path = self.template_dir / template_name
            if not template_path.exists():
                return {
                    'success': False,
                    'message': f'模板文件不存在：{template_name}'
                }
            
            # 加载模板
            tpl = DocxTemplate(str(template_path))
            
            # 渲染模板
            tpl.render(contract_data)
            
            # 生成文件名
            safe_contract_number = self.sanitize_filename(contract_data['contract_number'])
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            file_name = f'{safe_contract_number}_{timestamp}.docx'
            
            # 确保文件名安全
            file_name = self.sanitize_filename(file_name)
            
            # 保存文件
            file_path = self.generated_docs_dir / file_name
            tpl.save(str(file_path))
            
            # 插入地块图片到文档末尾
            self._insert_plot_images(str(file_path), contract_data['plots'])
            
            return {
                'success': True,
                'file_path': str(file_path),
                'file_name': file_name,
                'message': '合同生成成功'
            }
            
        except Exception as e:
            current_app.logger.error(f'生成合同文档失败：{str(e)}')
            return {
                'success': False,
                'message': f'生成失败：{str(e)}'
            }
    
    def download_contract(self, file_name):
        """
        下载合同文件
        
        Args:
            file_name: 文件名
            
        Returns:
            tuple: (file_path, download_name) 或 None
        """
        try:
            # 清理文件名
            safe_file_name = self.sanitize_filename(file_name)
            
            # 构建文件路径
            file_path = self.generated_docs_dir / safe_file_name
            
            # 检查文件是否存在
            if not file_path.exists():
                return None
            
            # 验证文件扩展名
            if not self.allowed_file(safe_file_name):
                return None
            
            return str(file_path), safe_file_name
            
        except Exception as e:
            current_app.logger.error(f'下载合同文件失败：{str(e)}')
            return None
    
    def cleanup_old_files(self, days=7):
        """
        清理过期的生成文件
        
        Args:
            days: 保留天数，默认 7 天
            
        Returns:
            int: 清理的文件数量
        """
        try:
            cleaned_count = 0
            cutoff_time = datetime.now().timestamp() - (days * 24 * 60 * 60)
            
            for file_path in self.generated_docs_dir.glob('*.docx'):
                if file_path.stat().st_mtime < cutoff_time:
                    try:
                        file_path.unlink()
                        cleaned_count += 1
                        current_app.logger.info(f'清理过期文件：{file_path.name}')
                    except Exception as e:
                        current_app.logger.error(f'清理文件失败：{file_path.name}, 错误：{str(e)}')
            
            return cleaned_count
            
        except Exception as e:
            current_app.logger.error(f'清理过期文件失败：{str(e)}')
            return 0
    
    def get_file_info(self, file_name):
        """
        获取文件信息
        
        Args:
            file_name: 文件名
            
        Returns:
            dict: 文件信息或 None
        """
        try:
            safe_file_name = self.sanitize_filename(file_name)
            file_path = self.generated_docs_dir / safe_file_name
            
            if not file_path.exists():
                return None
            
            stat = file_path.stat()
            return {
                'file_name': safe_file_name,
                'file_size': stat.st_size,
                'create_time': datetime.fromtimestamp(stat.st_ctime).strftime('%Y-%m-%d %H:%M:%S'),
                'modify_time': datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
            }
            
        except Exception as e:
            current_app.logger.error(f'获取文件信息失败：{str(e)}')
            return None


# 全局服务实例
contract_doc_service = ContractDocService()
